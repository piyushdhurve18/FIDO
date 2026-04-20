from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import util
from docx.enum.table import WD_TABLE_ALIGNMENT
import subprocess
import win32com.client

FILE_NAME = "FidoInternalTestReport"
HEADING_TEXT = "FIDO APPLET TEST SCENARIOS"
TESTER_NAME = "PIYUSH DHURVE"
DESCRIPTION = "ZTPASS FIDO APPLET CTAP v2.2"

SUMMARY_TABLE = None
DETAILED_TABLE = None

# def get_current_datetime():  ## In Format --> 2026-04-09 18:43:54
#     dt = datetime.now()
#     dt_str = dt.strftime("%Y-%m-%d %H:%M:%S")
#     return dt_str

def get_current_datetime():  ## In Format --> 09-04-2026 18:43:54 PM
    dt = datetime.now()
    dt_str = dt.strftime("%d-%b-%Y       %I:%M:%S %p")
    return dt_str

# def get_current_datetime():  ## In Format --> 09-Apr-2026 (06:42 PM)
#     dt = datetime.now()
#     return dt.strftime("%d-%b-%Y (%I:%M %p)")

def pct_width(pct):
    return Inches(6.7 * pct)   # 6.7 is total usable page width

def set_Heading(doc, headingText):
    """
    Creates:
    1. Centered Heading
    """
    # =========================
    # HEADING
    # =========================
    heading = doc.add_paragraph()
    run = heading.add_run(headingText)
    run.bold = True
    run.font.size = Pt(20)

    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Space after heading
    doc.add_paragraph("")



def create_header_section(doc, headingText, tester_name, description, datetime_value):
    """
    Creates:
    1. Centered Heading
    2. Header Table (30% / 70%)
    """

    set_Heading(doc, headingText)

    # =========================
    # TABLE
    # =========================
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER   # 👈 ADD THIS

    # Fix layout
    tbl = table._element
    tblPr = tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    table.columns[0].width = pct_width(0.3)
    table.columns[1].width = pct_width(0.7)

    labels = ["TESTER NAME", "DESCRIPTION", "DATE & TIME"]
    values = [tester_name, description, datetime_value]

    for i in range(3):
        # Set text
        table.rows[i].cells[0].text = labels[i]
        table.rows[i].cells[1].text = values[i]

        # 👉 Center align LEFT column
        for p in table.rows[i].cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 👉 Center align RIGHT column
        for p in table.rows[i].cells[1].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_table_borders(table)
    return table, tblPr

def add_summary_row(table, command, protocol, total, passed, failed):
    """
    Adds a new row to summary table with auto SR NO
    """

    # SR NO = current rows count (excluding header)
    sr_no = len(table.rows)

    row_cells = table.add_row().cells

    values = [sr_no, command, protocol, total, passed, failed]

    for i, val in enumerate(values):
        row_cells[i].text = str(val)

        # Center align (optional but recommended)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for p in row_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_summary_table(doc):
    """
    Creates summary table with headers:
    SR NO, COMMAND, PROTOCOL, TOTAL, PASSED, FAILED

    :param doc: Document object
    :param data_rows: List of rows (list of lists)
    """

    headers = [
        "SR NO.",
        "COMMAND",
        "PROTOCOL",
        "TOTAL NO. OF SCENARIOS",
        "NO. OF SCENARIOS PASSED",
        "NO. OF SCENARIOS FAILED"
    ]

    # =========================
    # CREATE TABLE
    # =========================
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER   # 👈 ADD THIS

    

    # Fix layout
    tbl = table._element
    tblPr = tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # =========================
    # SET COLUMN WIDTHS
    # (adjust for better PDF match)
    # =========================
    widths = [
        pct_width(0.08),
        pct_width(0.18),
        pct_width(0.14),
        pct_width(0.20),
        pct_width(0.20),
        pct_width(0.20),
    ]

    for i, width in enumerate(widths):
        table.columns[i].width = width

    # =========================
    # HEADER ROW
    # =========================

    header_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header

        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in paragraph.runs:
                run.bold = True

    # # =========================
    # # ADD DATA ROWS
    # # =========================
    # for row_data in data_rows:
    #     row_cells = table.add_row().cells
    #     for i, val in enumerate(row_data):
    #         row_cells[i].text = str(val)

    # =========================
    # BORDERS
    # =========================
    set_table_borders(table)

    return table, tblPr


def build_document(HEADING_TEXT, TESTER_NAME, DESCRIPTION):
    close_files()
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    create_header_section(doc, HEADING_TEXT, TESTER_NAME, DESCRIPTION, get_current_datetime())
    summaryTable, detailedTable = fetchTables(doc)
    return doc, summaryTable, detailedTable

def save_document(doc, docx_path):
    doc.save(docx_path)
    util.printcolor(util.BLUE, "Word File Created and Saved Successfully !!")


def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')   # solid line
        border.set(qn('w:sz'), '12')        # thickness (increase if needed)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') # black
        borders.append(border)

    tblPr.append(borders)


from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def create_detailed_table(doc):
    headers = ["SR NO.", "COMMAND", "PROTOCOL", "SCENARIO", "RESULT"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fix layout
    tbl = table._element
    tblPr = tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # Column widths (total ~15 cm like your summary table)
    widths = [
        Cm(1.2),   # SR NO
        Cm(3.0),   # COMMAND
        Cm(2.5),   # PROTOCOL
        Cm(6.3),   # SCENARIO (large)
        Cm(2.0),   # RESULT
    ]

    for i, width in enumerate(widths):
        table.columns[i].width = width

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header

        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    # Borders
    set_table_borders(table)

    return table


def add_detailed_row(table, command, protocol, scenario, result):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Auto SR NO
    sr_no = len(table.rows)

    row_cells = table.add_row().cells

    values = [sr_no, command, protocol, scenario, result]

    for i, val in enumerate(values):
        row_cells[i].text = str(val)

        # Alignment
        for p in row_cells[i].paragraphs:
            if i == 3:  # SCENARIO column (left align for readability)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

import os
from win32com.client import Dispatch

def convert_docx_to_pdf(input_path, output_path):
    word = Dispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=17)  # 17 = PDF
    doc.Close()
    word.Quit()
    util.printcolor(util.RED, "PDF File Created and Saved Successfully !!")


# def saveAllFiles(doc, FILE_NAME):   ### This function creating and saving docs in direct path
    # save_document(doc, FILE_NAME)

#     # Retrieving Paths
#     input_file = os.path.abspath(FILE_NAME + ".docx")
#     output_file = os.path.abspath(FILE_NAME + ".pdf")

#     convert_docx_to_pdf(input_file, output_file)



def saveAllFiles(doc, FILE_NAME): ### This function creating one folder 'REPORTS' and then creating and saving docs into it
    base_path = os.getcwd()
    reports_path = os.path.join(base_path, "REPORTS")

    os.makedirs(reports_path, exist_ok=True)

    docx_path = os.path.join(reports_path, FILE_NAME + ".docx")
    pdf_path = os.path.join(reports_path, FILE_NAME + ".pdf")

    # Save DOCX
    save_document(doc, docx_path)

    # Convert to PDF
    convert_docx_to_pdf(docx_path, pdf_path)

def fetchTables(doc):
    summaryTable, tbl = create_summary_table(doc)

    # 👉 Move next content to new page
    doc.add_page_break() 

    # Now anything added will go to next page
    detailedTable = create_detailed_table(doc)
    global SUMMARY_TABLE
    SUMMARY_TABLE = summaryTable

    global DETAILED_TABLE
    DETAILED_TABLE = detailedTable

    return summaryTable, detailedTable


def is_file_open(filepath):
    try:
        with open(filepath, "a"):
            return False
    except PermissionError:
        return True

def end_task(filepath):
    if is_file_open(filepath):
        util.printcolor(util.RED,"File is open → closing Word...")
    
    subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"], shell=True)

def close_word():
    
    subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"], shell=True)
    # util.printcolor(util.RED,"Closed all Word instances")



def close_pdf_readers():
    processes = [
        "AcroRd32.exe",   # Adobe Reader
        "Acrobat.exe",    # Adobe Pro
        "msedge.exe",     # Edge (PDF viewer)
        "chrome.exe",     # Chrome (PDF viewer)
        "FoxitReader.exe" # Foxit
    ]

    for proc in processes:
        subprocess.run(["taskkill", "/f", "/im", proc], shell=True)

    # util.printcolor(util.RED,"Closed all PDF reader instances")


def close_files():
    close_word()
    close_pdf_readers()
    util.printcolor(util.RED,"All Word & PDF Reader Instances Closed...")


    
